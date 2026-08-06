#!/usr/bin/env python
"""Build a visual Italian report from the privacy-checked real-field results."""

from __future__ import annotations

import argparse
import importlib.util
import json
import math
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from matplotlib.ticker import FuncFormatter

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_RESULTS = Path(r"C:\Users\Utente\Downloads\oncoemotion_real_results")
DEFAULT_OUT = ROOT / "reports" / "rapporto_risultati_oncoemotion_esmo.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "74A9CF"
TEAL = "2A9D8F"
GOLD = "C7922F"
RED = "B23A48"
GRAY = "6B7280"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "EAF2F8"
WHITE = "FFFFFF"
BLACK = "111827"

MODEL_LABELS = {
    "apertus-8b-instruct-2509": "Apertus 8B",
    "apertus-8b-meditronfo": "Apertus 8B MeditronFO",
    "eurollm-9b-instruct": "EuroLLM 9B",
    "eurollm-9b-meditronfo": "EuroLLM 9B MeditronFO",
    "gemma-3-27b-it": "Gemma 3 27B",
    "gemma-3-27b-meditronfo": "Gemma 3 27B MeditronFO",
    "ministral-8b-instruct-2410": "Ministral 8B",
    "qwen3-8b": "Qwen3 8B",
}


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _font(run, *, size=None, color=BLACK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _table_geometry(table, widths_dxa: list[int], *, indent_dxa=120) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"table widths must sum to 9360 DXA, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _cell_margins(cell)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_cell_text(cell, text, *, bold=False, color=BLACK, size=9, align="left") -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    _font(p.add_run(str(text)), size=size, color=color, bold=bold)


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    _font(run, size=8.5, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = _rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = _rgb(GRAY)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = False

    for header in (section.header, section.even_page_header):
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        _font(hp.add_run("ONCOEMOTION  |  Validazione su campi clinici reali"), size=8.5, color=GRAY)
    for footer in (section.footer, section.even_page_footer):
        fp = footer.paragraphs[0]
        _add_page_field(fp)


def _add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(5)
    _font(p.add_run("RAPPORTO DI RISULTATI"), size=9, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    _font(
        p.add_run("Rappresentazioni emozionali negli LLM\ne codifica dei sintomi oncologici"),
        size=24,
        color=NAVY,
        bold=True,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    _font(
        p.add_run(
            "Analisi del protocollo esmo-ai-2026-real-v2 su 1.275 valutazioni cliniche validate"
        ),
        size=12.5,
        color=GRAY,
    )

    metadata = [
        ("Data del rapporto", "6 agosto 2026"),
        ("Coorte", "8 modelli primari open-weight"),
        ("Disegno", "2 condizioni di ruolo, 20.400 osservazioni redatte"),
        ("Stato", "Run completo; controlli di integrità e privacy superati"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _font(p.add_run(f"{label}: "), size=9.5, color=GRAY, bold=True)
        _font(p.add_run(value), size=9.5, color=GRAY)


def _add_callout(doc: Document, label: str, value: str, detail: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _table_geometry(table, [9360])
    _repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    _shade(cell, PALE_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(label.upper()), size=9, color=BLUE, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(value), size=22, color=NAVY, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(detail), size=9.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    _font(p.add_run(text), size=11)


def _add_note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    _font(p.add_run(f"{label}: "), size=9.5, color=GOLD, bold=True)
    _font(p.add_run(text), size=9.5, color=GRAY)


def _add_figure(doc: Document, path: Path, caption: str, alt_text: str, *, width=6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    _font(p.add_run(f"Fonte: {text}"), size=8, color=GRAY, italic=True)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, font=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _table_geometry(table, widths)
    _repeat_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        _shade(cell, LIGHT_GRAY)
        _set_cell_text(cell, header, bold=True, color=NAVY, size=font, align="center")
    for row_values in rows:
        row = table.add_row()
        for index, (cell, value) in enumerate(zip(row.cells, row_values)):
            align = "left" if index == 0 else "center"
            _set_cell_text(cell, value, size=font, align=align)
    return table


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _save_figure(fig, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor=f"#{WHITE}")
    plt.close(fig)


def _style_axes(ax) -> None:
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="x", color="#D8DEE8", linewidth=0.8, alpha=0.8)
    ax.set_axisbelow(True)


def _comma_tick(value, _position):
    return f"{value:g}".replace(".", ",")


def _flow_chart(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 4.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    boxes = [
        (0.03, 0.56, 0.2, 0.25, "Dataset emozionale\nindipendente", "1.718 esempi\n33 concetti"),
        (0.27, 0.56, 0.2, 0.25, "Vettori per modello", "AUROC fuori campione\n+ gate lessicale"),
        (0.51, 0.56, 0.2, 0.25, "Campi clinici reali", "1.275 record validati\n63 item PRO"),
        (0.75, 0.56, 0.22, 0.25, "Misurazione LLM", "8 modelli x 2 ruoli\n20.400 righe"),
    ]
    for x, y, w, h, title, detail in boxes:
        rect = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.012,rounding_size=0.018",
            facecolor="#EAF2F8", edgecolor="#2E74B5", linewidth=1.5,
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + 0.16, title, ha="center", va="center", fontsize=11, weight="bold", color="#0B2545")
        ax.text(x + w / 2, y + 0.07, detail, ha="center", va="center", fontsize=9, color="#4B5563")
    for left, right in zip(boxes[:-1], boxes[1:]):
        ax.add_patch(FancyArrowPatch(
            (left[0] + left[2] + 0.008, 0.685), (right[0] - 0.008, 0.685),
            arrowstyle="-|>", mutation_scale=13, color="#6B7280", linewidth=1.2,
        ))
    summary = FancyBboxPatch(
        (0.19, 0.14), 0.62, 0.20, boxstyle="round,pad=0.015,rounding_size=0.018",
        facecolor="#F8FAFC", edgecolor="#C7922F", linewidth=1.5,
    )
    ax.add_patch(summary)
    ax.text(0.5, 0.265, "Analisi gerarchica", ha="center", va="center", fontsize=12, weight="bold", color="#0B2545")
    ax.text(
        0.5, 0.19,
        "Peso uguale per modello; bootstrap dei cluster source_id;\ncontrolli rigorosi di integrità e redazione",
        ha="center", va="center", fontsize=9.5, color="#4B5563",
    )
    ax.add_patch(FancyArrowPatch((0.86, 0.55), (0.67, 0.35), arrowstyle="-|>", mutation_scale=13, color="#6B7280", linewidth=1.2))
    _save_figure(fig, path)


def _forest_chart(report: dict, path: Path) -> None:
    affect = report["analysis"]["affect"]
    sensitivity = report["analysis"]["sensitivity"]
    entries = [
        ("Analisi primaria", affect["primary_error_association"]),
        ("Solo cluster non ambigui", sensitivity["unambiguous_identical_text_clusters"]["error_association"]),
        ("Testi con almeno 7 parole", sensitivity["texts_ge_7_words"]["error_association"]),
    ]
    fig, ax = plt.subplots(figsize=(9.5, 4.3))
    y = np.arange(len(entries))[::-1]
    colors = ["#2E74B5", "#74A9CF", "#C7922F"]
    for yi, (label, result), color in zip(y, entries, colors):
        estimate = result["estimate"] * 100
        lo, hi = [value * 100 for value in result["ci95"]]
        ax.errorbar(
            estimate, yi, xerr=[[estimate - lo], [hi - estimate]], fmt="o",
            color=color, ecolor=color, elinewidth=2, capsize=4, markersize=8,
        )
        ax.text(hi + 0.12, yi, f"{estimate:.2f}  [{lo:.2f}; {hi:.2f}]".replace(".", ","), va="center", fontsize=9, color="#374151")
    ax.axvline(0, color="#6B7280", linewidth=1.2)
    ax.set_yticks(y, [item[0] for item in entries])
    ax.set_xlabel("Variazione assoluta della probabilità di errore (punti percentuali)")
    ax.set_title("Associazione tra paura/preoccupazione ed errore", loc="left", color="#0B2545", weight="bold", fontsize=13)
    ax.xaxis.set_major_formatter(FuncFormatter(_comma_tick))
    ax.set_xlim(-3.2, 4.2)
    _style_axes(ax)
    fig.tight_layout()
    _save_figure(fig, path)


def _model_association_chart(model_assoc: dict[str, float], pooled: dict, path: Path) -> None:
    ordered = sorted(model_assoc, key=model_assoc.get)
    labels = [MODEL_LABELS[name] for name in ordered] + ["Stima aggregata"]
    values = [model_assoc[name] * 100 for name in ordered] + [pooled["estimate"] * 100]
    y = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(9.5, 5.6))
    colors = [f"#{RED}" if value < 0 else f"#{BLUE}" for value in values[:-1]]
    ax.scatter(values[:-1], y[:-1], s=68, c=colors, zorder=3)
    for yi, value in zip(y[:-1], values[:-1]):
        ax.plot([0, value], [yi, yi], color="#CBD5E1", linewidth=2, zorder=1)
        ax.text(value + (0.11 if value >= 0 else -0.11), yi, f"{value:.2f}".replace(".", ","), ha="left" if value >= 0 else "right", va="center", fontsize=8.5)
    estimate = values[-1]
    lo, hi = [value * 100 for value in pooled["ci95"]]
    ax.errorbar(estimate, y[-1], xerr=[[estimate - lo], [hi - estimate]], fmt="D", color=f"#{NAVY}", ecolor=f"#{NAVY}", capsize=4, markersize=7, zorder=4)
    ax.axvline(0, color="#6B7280", linewidth=1)
    ax.set_yticks(y, labels)
    ax.set_xlabel("Punti percentuali di errore per 1 DS del composito")
    ax.set_title("Eterogeneità descrittiva tra modelli", loc="left", color="#0B2545", weight="bold", fontsize=13)
    ax.xaxis.set_major_formatter(FuncFormatter(_comma_tick))
    ax.set_xlim(-1.6, 4.3)
    _style_axes(ax)
    fig.tight_layout()
    _save_figure(fig, path)


def _performance_chart(per_model: dict, path: Path) -> None:
    role = "oncologo"
    ordered = sorted(per_model, key=lambda name: per_model[name][role]["term_accuracy"])
    labels = [MODEL_LABELS[name] for name in ordered]
    accuracy = [per_model[name][role]["term_accuracy"] * 100 for name in ordered]
    macro = [per_model[name][role]["term_macro_recall"] * 100 for name in ordered]
    y = np.arange(len(ordered))
    fig, ax = plt.subplots(figsize=(9.5, 6.0))
    h = 0.35
    ax.barh(y - h / 2, accuracy, height=h, color="#2E74B5", label="Accuratezza")
    ax.barh(y + h / 2, macro, height=h, color="#74A9CF", label="Macro-recall")
    for yi, value in zip(y - h / 2, accuracy):
        ax.text(value + 0.6, yi, f"{value:.1f}".replace(".", ","), va="center", fontsize=8)
    for yi, value in zip(y + h / 2, macro):
        ax.text(value + 0.6, yi, f"{value:.1f}".replace(".", ","), va="center", fontsize=8)
    ax.set_yticks(y, labels)
    ax.set_xlabel("Percentuale")
    ax.set_xlim(0, 53)
    ax.set_title("Prestazione di codifica PRO-CTCAE", loc="left", color="#0B2545", weight="bold", fontsize=13)
    ax.legend(frameon=False, loc="lower right")
    ax.xaxis.set_major_formatter(FuncFormatter(_comma_tick))
    _style_axes(ax)
    fig.tight_layout()
    _save_figure(fig, path)


def _nonpro_chart(per_model: dict, path: Path) -> None:
    role = "oncologo"
    ordered = sorted(per_model, key=lambda name: per_model[name][role]["false_positive_rate_non_pro"])
    labels = [MODEL_LABELS[name] for name in ordered]
    abstain = [per_model[name][role]["abstention_rate_non_pro"] * 100 for name in ordered]
    false_pos = [per_model[name][role]["false_positive_rate_non_pro"] * 100 for name in ordered]
    y = np.arange(len(ordered))
    fig, ax = plt.subplots(figsize=(9.5, 5.7))
    ax.scatter(abstain, y, s=72, color=f"#{TEAL}", label="Astensione esplicita", zorder=3)
    ax.scatter(false_pos, y, s=72, color=f"#{RED}", label="Falso codice PRO", zorder=3)
    for yi, left, right in zip(y, abstain, false_pos):
        ax.plot([left, right], [yi, yi], color="#CBD5E1", linewidth=2, zorder=1)
    ax.set_yticks(y, labels)
    ax.set_xlabel("Percentuale dei 307 record non-PRO")
    ax.set_xlim(0, 65)
    ax.set_title("Comportamento sui casi che richiedevano astensione", loc="left", color="#0B2545", weight="bold", fontsize=13)
    ax.legend(frameon=False, loc="lower right")
    ax.xaxis.set_major_formatter(FuncFormatter(_comma_tick))
    _style_axes(ax)
    fig.tight_layout()
    _save_figure(fig, path)


def _vector_quality_chart(vector_quality: dict[str, dict[str, float]], path: Path) -> None:
    order = list(MODEL_LABELS)
    cols = ["Paura", "Preoccupazione", "Severità", "Valenza negativa"]
    fields = ["fear", "concern", "severity", "negval"]
    values = np.asarray([[vector_quality[name][field] for field in fields] for name in order])
    fig, ax = plt.subplots(figsize=(9.5, 5.9))
    image = ax.imshow(values, cmap="RdYlGn", vmin=0.45, vmax=0.85, aspect="auto")
    ax.set_xticks(np.arange(len(cols)), cols)
    ax.set_yticks(np.arange(len(order)), [MODEL_LABELS[name] for name in order])
    for i in range(values.shape[0]):
        for j in range(values.shape[1]):
            value = values[i, j]
            marker = "*" if value < 0.60 else ""
            ax.text(j, i, f"{value:.2f}{marker}".replace(".", ","), ha="center", va="center", color=f"#{WHITE}" if value < 0.53 or value > 0.78 else f"#{BLACK}", fontsize=9, weight="bold" if j < 2 else "normal")
    ax.add_patch(plt.Rectangle((-0.49, -0.49), 1.98, len(order) - 0.02, fill=False, edgecolor=f"#{BLUE}", linewidth=2))
    ax.set_title("Qualità fuori campione delle direzioni rappresentazionali", loc="left", color="#0B2545", weight="bold", fontsize=13)
    cbar = fig.colorbar(image, ax=ax, fraction=0.025, pad=0.03)
    cbar.set_label("AUROC")
    fig.text(0.01, 0.01, "* AUROC < 0,60. Il riquadro evidenzia i due assi primari.", fontsize=8.5, color="#6B7280")
    fig.tight_layout(rect=[0, 0.035, 1, 1])
    _save_figure(fig, path)


def _medicalization_chart(per_model: dict, path: Path) -> None:
    pairs = [
        ("Apertus 8B", "apertus-8b-instruct-2509", "apertus-8b-meditronfo"),
        ("EuroLLM 9B", "eurollm-9b-instruct", "eurollm-9b-meditronfo"),
        ("Gemma 3 27B", "gemma-3-27b-it", "gemma-3-27b-meditronfo"),
    ]
    fig, ax = plt.subplots(figsize=(8.8, 4.8))
    for index, (family, base, med) in enumerate(pairs):
        a = per_model[base]["oncologo"]["term_accuracy"] * 100
        b = per_model[med]["oncologo"]["term_accuracy"] * 100
        ax.plot([0, 1], [a, b], marker="o", markersize=8, linewidth=2.2, color=f"#{[BLUE, LIGHT_BLUE, TEAL][index]}", label=family)
        ax.text(-0.035, a, f"{a:.1f}".replace(".", ","), ha="right", va="center", fontsize=8.5)
        ax.text(1.035, b, f"{b:.1f}".replace(".", ","), ha="left", va="center", fontsize=8.5)
    ax.set_xticks([0, 1], ["Modello base", "Modello medicalizzato"])
    ax.set_ylabel("Accuratezza PRO-CTCAE (%)")
    ax.set_xlim(-0.2, 1.2)
    ax.set_ylim(22, 41)
    ax.set_title("Confronti base vs medicalizzato", loc="left", color="#0B2545", weight="bold", fontsize=13)
    ax.legend(frameon=False, loc="lower right")
    ax.yaxis.set_major_formatter(FuncFormatter(_comma_tick))
    ax.grid(axis="y", color="#D8DEE8", linewidth=0.8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    _save_figure(fig, path)


def _load_analyzer():
    path = ROOT / "scripts" / "analyze_real_fields.py"
    spec = importlib.util.spec_from_file_location("real_analysis_report", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _derived_metrics(results: Path, report: dict):
    analyzer = _load_analyzer()
    rows = []
    for path in sorted((results / "outputs" / "real_fields").glob("*__rows.jsonl")):
        rows.extend(analyzer._load_rows(path, path.name.split("__")[0]))
    analyzer._add_affect_composite(rows, "z_read", ["afraid_alarmed", "concerned"])
    model_assoc = {}
    for model in sorted({row["model"] for row in rows}):
        selected = [
            row for row in rows
            if row["model"] == model and row["role"] == "oncologo"
            and row["arm"] == "intact" and row["gold_class"] == "term"
        ]
        model_assoc[model] = analyzer._error_affect_slope(selected)

    term = [
        row for row in rows
        if row["role"] == "oncologo" and row["arm"] == "intact"
        and row["gold_class"] == "term"
    ]
    length = lambda row: math.log1p(float(row.get("n_words") or 0))
    def z(row, name, control=False):
        return analyzer._z_value(row, "z_read", name, control=control)
    def slope(nuisances):
        values = []
        for model in sorted({row["model"] for row in term}):
            selected = [row for row in term if row["model"] == model]
            values.append(analyzer._residual_association(
                selected,
                exposure=lambda row: row.get("_affect_composite"),
                outcome=lambda row: float(not bool(row["correct"])),
                group=lambda row: (row.get("source_item"), row.get("grade")),
                nuisances=nuisances,
            ))
        return float(np.mean(values)), sum(value > 0 for value in values)
    no_cov, no_cov_positive = slope([])
    length_only, length_positive = slope([length])
    all_controls, all_positive = slope([
        length,
        lambda row: z(row, "clinical_severity", True),
        lambda row: z(row, "general_negative_valence", True),
    ])
    return rows, model_assoc, {
        "no_cov": no_cov,
        "no_cov_positive": no_cov_positive,
        "length_only": length_only,
        "length_positive": length_positive,
        "all_controls": all_controls,
        "all_positive": all_positive,
    }


def _vector_metrics(results: Path) -> dict[str, dict[str, float]]:
    output = {}
    for path in sorted((results / "outputs" / "models_real").glob("*/vector_validation.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        gate = data["protocol_gate"]["per_axis"]
        concepts = data["concepts"]
        output[path.parent.name] = {
            "fear": gate["afraid_alarmed"]["auroc"],
            "concern": gate["concerned"]["auroc"],
            "severity": concepts["clinical_severity"]["best_auroc"],
            "negval": concepts["general_negative_valence"]["best_auroc"],
            "fear_lex": gate["afraid_alarmed"]["max_abs_lexical_cosine"],
            "concern_lex": gate["concerned"]["max_abs_lexical_cosine"],
        }
    return output


def _fmt_pct(value: float, digits=1) -> str:
    return f"{value * 100:.{digits}f}%".replace(".", ",")


def _fmt_pp(value: float, digits=2) -> str:
    return f"{value * 100:+.{digits}f} pp".replace(".", ",")


def build_report(results: Path, out: Path) -> list[Path]:
    analysis_path = results / "outputs" / "real_fields" / "real_primary_analysis.json"
    report = json.loads(analysis_path.read_text(encoding="utf-8"))
    analysis = report["analysis"]
    per_model = analysis["per_model_role"]
    pooled = analysis["pooled_primary_role"]
    affect = analysis["affect"]
    sensitivity = analysis["sensitivity"]
    _, model_assoc, robustness = _derived_metrics(results, report)
    vector_quality = _vector_metrics(results)

    assets = out.parent / "assets"
    figures = {
        "flow": assets / "01_study_flow.png",
        "forest": assets / "02_primary_forest.png",
        "association": assets / "03_model_association.png",
        "performance": assets / "04_model_performance.png",
        "nonpro": assets / "05_nonpro_behavior.png",
        "vectors": assets / "06_vector_quality.png",
        "medical": assets / "07_medicalization.png",
    }
    _flow_chart(figures["flow"])
    _forest_chart(report, figures["forest"])
    _model_association_chart(model_assoc, affect["primary_error_association"], figures["association"])
    _performance_chart(per_model, figures["performance"])
    _nonpro_chart(per_model, figures["nonpro"])
    _vector_quality_chart(vector_quality, figures["vectors"])
    _medicalization_chart(per_model, figures["medical"])

    doc = Document()
    _configure_document(doc)
    doc.core_properties.title = "Rappresentazioni emozionali negli LLM e codifica dei sintomi oncologici"
    doc.core_properties.subject = "Rapporto analitico del protocollo esmo-ai-2026-real-v2"
    doc.core_properties.keywords = "oncologia, LLM, emozioni, PRO-CTCAE, ESMO"
    _add_title_block(doc)
    doc.add_paragraph()
    primary = affect["primary_error_association"]
    lo, hi = primary["ci95"]
    _add_callout(
        doc,
        "Risultato principale",
        _fmt_pp(primary["estimate"]),
        f"di probabilità d'errore per 1 DS di paura/preoccupazione; IC 95% da {_fmt_pp(lo)} a {_fmt_pp(hi)}",
    )
    doc.add_heading("Sintesi esecutiva", level=1)
    _add_bullet(
        doc,
        "Il run è completo e tecnicamente valido: 8 modelli, 20.400 osservazioni, tutti i controlli di coorte, hash e redazione superati.",
    )
    _add_bullet(
        doc,
        "Le direzioni indipendenti di paura e preoccupazione superano il gate in tutti i modelli e mostrano una piccola associazione osservazionale con l'errore di codifica.",
    )
    _add_bullet(
        doc,
        "Il prompt da oncologo cambia il codice nel 14,8% dei casi, ma non determina uno spostamento affettivo significativo e non migliora descrittivamente l'accuratezza.",
    )
    _add_note(
        doc,
        "Lettura corretta",
        "Le proiezioni misurano geometria rappresentazionale. Non dimostrano che il modello provi emozioni e, sui testi reali, non consentono una conclusione causale.",
    )

    _page_break(doc)
    doc.add_heading("1. Che cosa è stato fatto", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Lo studio ha separato deliberatamente la costruzione delle direzioni emozionali dalla loro verifica sui dati clinici. "
        "Il dataset sintetico non stabilisce la correttezza delle associazioni cliniche: serve soltanto a imparare, per ogni modello, assi rappresentazionali di emozione. "
        "Il file clinico validato fornisce invece il gold standard per misurare codifica, astensione ed errore."
    )
    _add_figure(
        doc,
        figures["flow"],
        "Figura 1. Flusso dello studio: apprendimento indipendente degli assi, applicazione al testo reale e analisi gerarchica.",
        "Diagramma del flusso: 1.718 esempi sintetici, vettori per modello, 1.275 record clinici, 8 modelli per 2 ruoli e analisi gerarchica.",
    )
    _add_source(doc, "protocollo esmo-ai-2026-real-v2 e manifest del pacchetto dei risultati.")
    doc.add_heading("Integrità del run", level=2)
    integrity = report["artifact_validation"]
    source = integrity["source_counts"]
    _add_table(
        doc,
        ["Controllo", "Esito"],
        [
            ["Coorte", "8/8 modelli primari presenti"],
            ["Osservazioni", f"{analysis['n_rows']:,}".replace(",", ".") + " righe complete"],
            ["Fonte clinica", f"{source['records']:,}".replace(",", ".") + " record; 63 item PRO"],
            ["Composizione", "968 record PRO; 245 CTCAE senza target PRO; 62 non associabili"],
            ["Testi distinti", f"{source['unique_source_ids']:,}".replace(",", ".") + " source_id non reversibili"],
            ["Privacy", "Testo e generazioni grezze assenti dal pacchetto"],
            ["Manifest", "36/36 file presenti con SHA-256 corrispondente"],
        ],
        [2700, 6660],
        font=9.2,
    )

    _page_break(doc)
    doc.add_heading("2. Risultato principale", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Nel ruolo predefinito di oncologo, una deviazione standard in più del composito paura/preoccupazione è associata a "
    )
    _font(p.add_run(_fmt_pp(primary["estimate"])), bold=True, color=NAVY)
    p.add_run(
        " di probabilità assoluta d'errore, dopo il centraggio entro item e grado e l'aggiustamento per lunghezza, severità clinica e valenza negativa. "
        "L'intervallo di confidenza è interamente sopra zero, ma il margine inferiore è molto vicino allo zero."
    )
    _add_figure(
        doc,
        figures["forest"],
        "Figura 2. Stima primaria e analisi di sensibilità. Le barre indicano gli IC 95% del bootstrap gerarchico.",
        "Forest plot dell'associazione errore-paura/preoccupazione: stima primaria positiva, sensibilità positiva nei cluster non ambigui, risultato incerto nei testi lunghi.",
    )
    _add_source(doc, "real_primary_analysis.json; 5.000 ricampionamenti, peso uguale per modello.")
    doc.add_heading("Come interpretarlo", level=2)
    _add_bullet(doc, "L'effetto è piccolo: circa 1-2 errori aggiuntivi ogni 100 valutazioni per 1 DS del composito.")
    _add_bullet(doc, "La sensibilità che esclude i testi identici con annotazioni discordanti resta positiva: " + _fmt_pp(sensitivity["unambiguous_identical_text_clusters"]["error_association"]["estimate"]) + ".")
    _add_bullet(doc, "Nei testi con almeno 7 parole la stima è vicina allo zero e molto imprecisa; questo sottogruppo contiene solo 856 righe modello-record.")
    _add_note(doc, "Forza dell'evidenza", "Il risultato sostiene un segnale osservazionale replicato, non una relazione causale e non un effetto clinicamente grande.")

    _page_break(doc)
    doc.add_heading("3. Quanto è uniforme tra i modelli?", level=1)
    p = doc.add_paragraph(
        "L'associazione aggiustata è positiva in sei modelli su otto. EuroLLM e Gemma mostrano le stime più alte; Ministral e Qwen mostrano stime lievemente negative. "
        "Il bootstrap aggregato incorpora questa eterogeneità attribuendo lo stesso peso a ogni modello."
    )
    _add_figure(
        doc,
        figures["association"],
        "Figura 3. Stime puntuali descrittive per modello; la riga aggregata riporta anche l'IC 95%.",
        "Grafico per modello dell'aumento o riduzione della probabilità di errore associata al composito paura-preoccupazione.",
    )
    _add_source(doc, "ricalcolo sui soli output redatti; gli IC per modello non erano endpoint predefiniti.")
    doc.add_heading("Robustezza ai controlli", level=2)
    _add_table(
        doc,
        ["Specifica descrittiva", "Stima", "Modelli positivi"],
        [
            ["Senza covariate", _fmt_pp(robustness["no_cov"]), f"{robustness['no_cov_positive']}/8"],
            ["Solo lunghezza", _fmt_pp(robustness["length_only"]), f"{robustness['length_positive']}/8"],
            ["Tutti i controlli predefiniti", _fmt_pp(robustness["all_controls"]), f"{robustness['all_positive']}/8"],
        ],
        [5200, 2080, 2080],
        font=9.2,
    )
    _add_note(doc, "Punto metodologico", "Le direzioni di severità e valenza negativa sono deboli in diversi modelli. Tuttavia, la stima con la sola lunghezza (1,37 pp) è quasi identica a quella pienamente aggiustata (1,42 pp).")

    _page_break(doc)
    doc.add_heading("4. Prestazione di codifica", level=1)
    p = doc.add_paragraph(
        "Il compito resta difficile. Nel ruolo oncologo, l'accuratezza aggregata è 31,8% e la macro-recall è 39,3%. "
        "La macro-recall più alta dell'accuratezza riflette il peso uguale assegnato ai 63 item, indipendentemente dalla loro frequenza."
    )
    _add_figure(
        doc,
        figures["performance"],
        "Figura 4. Accuratezza e macro-recall sui 968 record con gold PRO-CTCAE, ruolo oncologo.",
        "Barre orizzontali di accuratezza e macro-recall per gli otto modelli; Gemma MeditronFO è il migliore, Qwen il più basso.",
    )
    _add_source(doc, "per_model_role del rapporto primario.")

    table_rows = []
    for model in sorted(per_model, key=lambda name: per_model[name]["oncologo"]["term_accuracy"], reverse=True):
        data = per_model[model]["oncologo"]
        table_rows.append([
            MODEL_LABELS[model],
            _fmt_pct(data["term_accuracy"]),
            _fmt_pct(data["term_macro_recall"]),
            _fmt_pct(data["abstention_rate_non_pro"]),
            _fmt_pct(data["false_positive_rate_non_pro"]),
            _fmt_pp(model_assoc[model]),
        ])
    _add_table(
        doc,
        ["Modello", "Acc.", "Macro-rec.", "Astens.", "Falsi PRO", "Assoc. errore"],
        table_rows,
        [3600, 1037, 1181, 1080, 1152, 1310],
        font=8.2,
    )

    _page_break(doc)
    doc.add_heading("5. Il problema dell'astensione", level=1)
    p = doc.add_paragraph(
        "I 307 record senza un target PRO diretto erano il banco di prova dell'astensione. In media, i modelli si astengono esplicitamente solo nel 12,5% dei casi e producono un falso codice PRO nel 50,3%. "
        "La parte restante comprende risposte non mappabili o non classificabili come astensione esplicita."
    )
    _add_figure(
        doc,
        figures["nonpro"],
        "Figura 5. Astensione esplicita e falso codice PRO sui record non-PRO, ruolo oncologo.",
        "Confronto per modello tra percentuale di astensione esplicita e percentuale di falso codice PRO sui casi non-PRO.",
    )
    _add_source(doc, "307 record per modello nel ruolo oncologo.")
    _add_note(doc, "Trade-off", "Apertus base ha la minore frequenza di falso codice PRO (27,4%), ma anche una delle accuratezze più basse sui record PRO. Gli altri modelli si concentrano intorno al 50-57% di falsi positivi non-PRO.")
    doc.add_heading("Effetto del ruolo", level=2)
    p = doc.add_paragraph()
    p.add_run("Il prompt oncologo e il controllo filler scelgono codici differenti nel ")
    _font(p.add_run(_fmt_pct(pooled["role_top1_disagreement"]["estimate"])), bold=True, color=NAVY)
    p.add_run(" dei record PRO (IC 95% 11,7%-18,4%). Tuttavia, lo spostamento del composito affettivo è ")
    shift = affect["role_shift_affect_composite"]
    _font(p.add_run(f"{shift['estimate']:+.3f} DS".replace(".", ",")), bold=True, color=NAVY)
    p.add_run(f" con IC 95% [{shift['ci95'][0]:+.3f}; {shift['ci95'][1]:+.3f}]".replace(".", ",") + ", quindi compatibile con nessun effetto.")

    _page_break(doc)
    doc.add_heading("6. Gli assi emozionali sono validi?", level=1)
    p = doc.add_paragraph(
        "Sì per l'endpoint primario. Paura e preoccupazione superano in tutti gli otto modelli i tre criteri predefiniti: almeno 60 esempi positivi fuori campione, AUROC almeno 0,60 e bassa somiglianza con gli assi di semplice menzione o negazione lessicale."
    )
    _add_figure(
        doc,
        figures["vectors"],
        "Figura 6. AUROC fuori campione degli assi primari e delle due principali covariate rappresentazionali.",
        "Mappa di calore AUROC per paura, preoccupazione, severità clinica e valenza negativa negli otto modelli. Gli assi primari superano 0,60; molte covariate no.",
    )
    _add_source(doc, "vector_validation.json di ciascun modello; split per famiglia di parafrasi.")
    fear_range = (min(value["fear"] for value in vector_quality.values()), max(value["fear"] for value in vector_quality.values()))
    concern_range = (min(value["concern"] for value in vector_quality.values()), max(value["concern"] for value in vector_quality.values()))
    max_lex = max(max(value["fear_lex"], value["concern_lex"]) for value in vector_quality.values())
    _add_bullet(doc, f"Paura: AUROC {fear_range[0]:.2f}-{fear_range[1]:.2f}; preoccupazione: {concern_range[0]:.2f}-{concern_range[1]:.2f}.".replace(".", ","))
    _add_bullet(doc, f"Massima correlazione assoluta con i controlli lessicali: {max_lex:.3f}, ampiamente sotto la soglia 0,50.".replace(".", ","))
    _add_bullet(doc, "Severità clinica e valenza negativa sono meno separabili, soprattutto nei modelli Apertus, Gemma e Qwen; per questo vengono trattate come covariate imperfette, non come endpoint.")

    _page_break(doc)
    doc.add_heading("7. Medicalizzazione e confronto delle famiglie", level=1)
    p = doc.add_paragraph(
        "Nei tre confronti appaiati disponibili, le versioni MeditronFO hanno un'accuratezza maggiore delle rispettive basi. La differenza media descrittiva è di circa 3,3 punti percentuali, ma il protocollo non ha prodotto un intervallo inferenziale specifico per questo confronto."
    )
    _add_figure(
        doc,
        figures["medical"],
        "Figura 7. Accuratezza dei modelli base e delle corrispondenti versioni medicalizzate.",
        "Grafico appaiato delle accuratezze base e medicalizzate per Apertus, EuroLLM e Gemma; tutte le linee aumentano.",
    )
    _add_source(doc, "accuratezza nel ruolo oncologo; confronto descrittivo.")
    pair_rows = []
    for family, base, med in [
        ("Apertus 8B", "apertus-8b-instruct-2509", "apertus-8b-meditronfo"),
        ("EuroLLM 9B", "eurollm-9b-instruct", "eurollm-9b-meditronfo"),
        ("Gemma 3 27B", "gemma-3-27b-it", "gemma-3-27b-meditronfo"),
    ]:
        a = per_model[base]["oncologo"]["term_accuracy"]
        b = per_model[med]["oncologo"]["term_accuracy"]
        pair_rows.append([family, _fmt_pct(a), _fmt_pct(b), _fmt_pp(b - a)])
    _add_table(doc, ["Famiglia", "Base", "Medicalizzato", "Differenza"], pair_rows, [3744, 1872, 1872, 1872], font=9.2)
    _add_note(doc, "Interpretazione", "La medicalizzazione sembra migliorare il comportamento di codifica, ma non rende uniforme l'associazione emozione-errore e non risolve il problema dell'astensione sui casi non-PRO.")

    _page_break(doc)
    doc.add_heading("8. Che cosa si può concludere", level=1)
    doc.add_heading("Messaggio scientifico sostenibile", level=2)
    p = doc.add_paragraph()
    _font(p.add_run(
        "Direzioni di paura e preoccupazione apprese indipendentemente generalizzano a testo oncologico reale validato e mostrano una piccola associazione con l'errore di codifica PRO-CTCAE, mentre il framing da oncologo modifica le decisioni senza uno spostamento affettivo robusto."
    ), bold=True, color=NAVY)
    doc.add_heading("Cosa non bisogna affermare", level=2)
    _add_bullet(doc, "Non dire che gli LLM provano paura o preoccupazione: si misurano proiezioni nello spazio delle rappresentazioni interne.")
    _add_bullet(doc, "Non dire che paura e preoccupazione causano l'errore: i campi reali non hanno un controfattuale neutro appaiato.")
    _add_bullet(doc, "Non presentare il risultato come forte o definitivo: l'effetto è piccolo, l'IC sfiora lo zero e due modelli mostrano una direzione opposta.")
    _add_bullet(doc, "Non generalizzare ai testi narrativi lunghi: quella sensibilità è sottodimensionata e non conclusiva.")
    doc.add_heading("Perché è rilevante per ESMO", level=2)
    p = doc.add_paragraph(
        "Il contributo non è una nuova ontologia delle emozioni, ma una verifica oncologica concreta: la geometria affettiva interna degli LLM è misurabile su campi sintomatologici reali, è parzialmente associata alla sicurezza della codifica e interagisce con differenze di modello e medicalizzazione. "
        "Il risultato più clinicamente evidente resta però la difficoltà di astensione, che produce molti falsi codici PRO nei casi senza target appropriato."
    )
    doc.add_heading("Priorità analitiche successive", level=2)
    _add_bullet(doc, "Aggiungere intervalli per i confronti base-medicalizzato e per le stime di ciascun modello.")
    _add_bullet(doc, "Separare paura e preoccupazione come endpoint esplorativi, mantenendo il composito come endpoint primario.")
    _add_bullet(doc, "Ripetere l'analisi su un campione esterno con più testi narrativi di almeno 7 parole.")
    _add_bullet(doc, "Analizzare errori per item PRO e categorie di risposta, senza esporre testo clinico grezzo.")

    _page_break(doc)
    doc.add_heading("Appendice tecnica", level=1)
    _add_table(
        doc,
        ["Elemento", "Definizione operativa"],
        [
            ["Endpoint primario", "Variazione assoluta della probabilità di errore per 1 DS entro modello del composito paura/preoccupazione."],
            ["Popolazione primaria", "968 record con associazione PRO-CTCAE, ruolo oncologo, braccio intact."],
            ["Controllo di gruppo", "Centraggio entro item sorgente e grado associato."],
            ["Covariate", "Log-lunghezza, proiezione di severità clinica e proiezione di valenza negativa."],
            ["Inferenza", "Bootstrap gerarchico dei source_id entro modello, ricampionamento dei modelli, peso uguale per modello, 5.000 draw."],
            ["Gate affettivo", "AUROC fuori campione >=0,60; almeno 60 positivi; massimo |coseno lessicale| <=0,50."],
            ["Privacy", "Nel pacchetto esportato text=source_id; model_generated nullo; workbook, attivazioni e terminologia ufficiale esclusi."],
            ["Interpretazione", "Associazione osservazionale di una rappresentazione interna; nessuna attribuzione di esperienza soggettiva."],
        ],
        [2700, 6660],
        font=8.8,
    )
    _add_source(
        doc,
        "oncoemotion_real_results: real_primary_analysis.json, vector_validation.json, meta e manifest; protocollo esmo-ai-2026-real-v2.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    _font(p.add_run("Nota di stato"), size=9, color=BLUE, bold=True)
    p = doc.add_paragraph()
    _font(
        p.add_run(
            "Questo documento descrive e interpreta risultati già prodotti. Non costituisce abstract, poster, manoscritto clinico né validazione normativa del software."
        ),
        size=9,
        color=GRAY,
        italic=True,
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return list(figures.values())


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--results", type=Path, default=DEFAULT_RESULTS)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    figures = build_report(args.results, args.out)
    print(f"Wrote {args.out}")
    print(f"Figures: {len(figures)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
